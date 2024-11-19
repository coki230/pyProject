import spacy
import pathlib
import docx
import os


# Load English tokenizer, tagger, parser and NER
nlp = spacy.load("en_core_web_sm")

# Process whole documents
text = ("When Sebastian Thrun started working on self-driving cars at "
        "Google in 2007, few people outside of the company took him "
        "seriously. “I can tell you very senior CEOs of major American "
        "car companies would shake my hand and turn away because I wasn’t "
        "worth talking to,” said Thrun, in an interview with Recode earlier "
        "this week.")
doc = nlp(text)

# Analyze syntax
print("Noun phrases:", [chunk.text for chunk in doc.noun_chunks])
print("Verbs:", [token.lemma_ for token in doc if token.pos_ == "VERB"])

# Find named entities, phrases and concepts
for entity in doc.ents:
    print(entity.text, entity.label_)


# directory = "C:\\"
# files = list(pathlib.Path(directory).glob("**/*.docx"))
# print(files)



# # Load the .docx file
# file_path = '33.docx'
# document = docx.Document(file_path)
#
# # Define the keywords you want to search for
# keywords = ['document.add_heading']
#
# # Initialize a list to store the matching text
# matches = []
#
# # Iterate through each paragraph in the document
# for para in document.paragraphs:
#     for keyword in keywords:
#         if keyword.lower() in para.text.lower():
#             matches.append(para.text)
#
# # tables = document.tables
# # for table in tables:
# #     for row in table.rows:
# #         for cell in row.cells:
# #             print(cell.text)
# #             for keyword in keywords:
# #                 if keyword.lower() in cell.text.lower():
# #                     matches.append(cell.text)
#
# print(matches)


os.system("33.docx")